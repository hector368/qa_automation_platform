"""Pruebas de las vistas del generador de casos de prueba."""

from __future__ import annotations
from openpyxl import load_workbook

from apps.test_cases.services.ado_rows_builder import (
    ADO_NCOLS,
)
from apps.test_cases.services.xlsx_generator import (
    generate_xlsx,
)
import json
from io import BytesIO
from unittest.mock import Mock, patch

from django.core.cache import cache
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase, override_settings
from django.urls import reverse
from docx import Document

from apps.test_cases.services.orchestrator import (
    PreparedGeneration,
)
from apps.test_cases.services.requirement_splitter import (
    RequirementBlock,
)


@override_settings(
    TEST_CASES_RESULT_TTL_SECONDS=60,
)
class TestCasesViewsTests(TestCase):
    """Pruebas HTTP de análisis, generación y descarga."""

    def setUp(self) -> None:
        """Limpia el caché antes de cada prueba."""
        cache.clear()

    def tearDown(self) -> None:
        """Limpia el caché después de cada prueba."""
        cache.clear()

    def test_home_is_available(self) -> None:
        """Comprueba que la interfaz original esté disponible."""
        response = self.client.get(
            reverse("test_cases:home")
        )
    
        self.assertEqual(
            response.status_code,
            200,
        )
    
        self.assertTemplateUsed(
            response,
            "test_cases/index.html",
        )
    
        self.assertContains(
            response,
            "Automated Test Case Generator",
        )
    
        self.assertContains(
            response,
            "test_cases/css/test_cases.css",
        )
    
        self.assertContains(
            response,
            "test_cases/js/test_cases.js",
        )
    
        self.assertContains(
            response,
            "test_cases/img/b-logo.png",
        )
    
        self.assertContains(
            response,
            "test_cases/img/logo.png",
        )
    def test_analyzes_valid_docx(self) -> None:
        """Analiza un DOCX válido sin llamar a Claude."""
        uploaded_file = SimpleUploadedFile(
            name="CFC.003_PDD.docx",
            content=self._build_valid_docx(),
            content_type=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
        )

        response = self.client.post(
            reverse("test_cases:analyze"),
            {
                "document": uploaded_file,
            },
        )

        self.assertEqual(
            response.status_code,
            200,
        )

        payload = response.json()

        self.assertTrue(
            payload["ok"]
        )

        self.assertEqual(
            payload["project_id"],
            "CFC.003",
        )

        self.assertEqual(
            payload["total_blocks"],
            2,
        )

        self.assertEqual(
            len(payload["requirements"]),
            2,
        )

    def test_rejects_analysis_without_file(self) -> None:
        """Rechaza una solicitud de análisis sin documento."""
        response = self.client.post(
            reverse("test_cases:analyze")
        )

        self.assertEqual(
            response.status_code,
            400,
        )

        self.assertEqual(
            response.json()["code"],
            "ERR_NO_FILE",
        )

    def test_rejects_unsupported_extension(self) -> None:
        """Rechaza extensiones diferentes de PDF y DOCX."""
        uploaded_file = SimpleUploadedFile(
            name="document.txt",
            content=b"contenido",
            content_type="text/plain",
        )

        response = self.client.post(
            reverse("test_cases:analyze"),
            {
                "document": uploaded_file,
            },
        )

        self.assertEqual(
            response.status_code,
            400,
        )

        self.assertEqual(
            response.json()["code"],
            "ERR_BAD_EXT",
        )

    def test_returns_422_when_requirements_are_missing(
        self,
    ) -> None:
        """Devuelve 422 cuando no detecta requerimientos."""
        uploaded_file = SimpleUploadedFile(
            name="document.docx",
            content=self._build_docx_without_requirements(),
            content_type=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
        )

        response = self.client.post(
            reverse("test_cases:analyze"),
            {
                "document": uploaded_file,
            },
        )

        self.assertEqual(
            response.status_code,
            422,
        )

        self.assertEqual(
            response.json()["code"],
            "ERR_NO_REQUIREMENTS",
        )

    @patch(
        "apps.test_cases.views.generate_document_service"
    )
    def test_generates_and_stores_result(
        self,
        generate_mock: Mock,
    ) -> None:
        """Genera mediante JSON y guarda el XLSX temporalmente."""
        generate_mock.return_value = (
            self._build_generation_result()
        )

        uploaded_file = SimpleUploadedFile(
            name="CFC.003_PDD.docx",
            content=b"test-document",
            content_type=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
        )

        response = self.client.post(
            reverse("test_cases:generate"),
            {
                "document": uploaded_file,
                "assigned_to": "Usuario QA",
                "selected_requirements": "1",
            },
        )

        self.assertEqual(
            response.status_code,
            200,
        )

        payload = response.json()

        self.assertTrue(
            payload["ok"]
        )

        self.assertEqual(
            payload["code"],
            "OK_GENERATED",
        )

        self.assertIn(
            "result_id",
            payload,
        )

        self.assertIn(
            "download_url",
            payload,
        )

        self.assertEqual(
            payload["selected_requirements"],
            [1],
        )

        download_response = self.client.get(
            payload["download_url"]
        )

        self.assertEqual(
            download_response.status_code,
            200,
        )

        self.assertEqual(
            download_response["Content-Type"],
            (
                "application/vnd.openxmlformats-officedocument."
                "spreadsheetml.sheet"
            ),
        )

        self.assertIn(
            "CFC.003_PDD_TC.xlsx",
            download_response[
                "Content-Disposition"
            ],
        )

    def test_rejects_generation_without_file(self) -> None:
        """Rechaza la generación cuando falta el documento."""
        response = self.client.post(
            reverse("test_cases:generate"),
            {
                "assigned_to": "Usuario QA",
            },
        )

        self.assertEqual(
            response.status_code,
            400,
        )

        self.assertEqual(
            response.json()["code"],
            "ERR_NO_FILE",
        )

    def test_rejects_generation_without_assigned_to(
        self,
    ) -> None:
        """Rechaza la generación cuando Assigned To está vacío."""
        uploaded_file = SimpleUploadedFile(
            name="document.docx",
            content=b"test-document",
            content_type=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
        )

        response = self.client.post(
            reverse("test_cases:generate"),
            {
                "document": uploaded_file,
                "assigned_to": "",
            },
        )

        self.assertEqual(
            response.status_code,
            400,
        )

        self.assertEqual(
            response.json()["code"],
            "ERR_GENERATION_INPUT",
        )

    def test_rejects_invalid_requirement_selection(
        self,
    ) -> None:
        """Rechaza selecciones con formato inválido."""
        uploaded_file = SimpleUploadedFile(
            name="document.docx",
            content=b"test-document",
            content_type=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
        )

        response = self.client.post(
            reverse("test_cases:generate"),
            {
                "document": uploaded_file,
                "assigned_to": "Usuario QA",
                "selected_requirements": "1,dos,3",
            },
        )

        self.assertEqual(
            response.status_code,
            400,
        )

        self.assertEqual(
            response.json()["code"],
            "ERR_GENERATION_INPUT",
        )

    @patch(
        "apps.test_cases.views."
        "iter_generate_prepared_document"
    )
    @patch(
        "apps.test_cases.views."
        "prepare_generation_document"
    )
    def test_streams_generation_progress(
        self,
        prepare_mock: Mock,
        iter_generate_mock: Mock,
    ) -> None:
        """
        Transmite eventos NDJSON y almacena el resultado final.

        La prueba no consume Claude porque los servicios están simulados.
        """
        prepare_mock.return_value = PreparedGeneration(
            filename="CFC.003_PDD.docx",
            project_id="CFC.003",
            context_text="Sistema: SAP",
            blocks=(
                RequirementBlock(
                    requirement_number=1,
                    scenario_name="Validar archivo",
                    input_text="Validar archivo.",
                ),
            ),
        )

        generation_result = (
            self._build_generation_result()
        )

        iter_generate_mock.return_value = iter(
            [
                {
                    "type": "started",
                    "ok": True,
                    "project_id": "CFC.003",
                    "total_requirements": 1,
                    "progress": 0,
                },
                {
                    "type": "requirement_started",
                    "ok": True,
                    "requirement_number": 1,
                    "scenario_name": "Validar archivo",
                    "current": 1,
                    "total": 1,
                    "progress": 0,
                },
                {
                    "type": "requirement_completed",
                    "ok": True,
                    "requirement_number": 1,
                    "current": 1,
                    "total": 1,
                    "progress": 100,
                },
                {
                    "type": "completed",
                    "ok": True,
                    "progress": 100,
                    "result": generation_result,
                },
            ]
        )

        uploaded_file = SimpleUploadedFile(
            name="CFC.003_PDD.docx",
            content=b"document-content",
            content_type=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
        )

        response = self.client.post(
            reverse(
                "test_cases:generate_stream"
            ),
            {
                "document": uploaded_file,
                "assigned_to": "Usuario QA",
                "selected_requirements": "1",
            },
        )

        self.assertEqual(
            response.status_code,
            200,
        )

        self.assertTrue(
            response.streaming
        )

        self.assertEqual(
            response["Content-Type"],
            "application/x-ndjson; charset=utf-8",
        )

        raw_content = b"".join(
            response.streaming_content
        ).decode("utf-8")

        events = [
            json.loads(line)
            for line in raw_content.splitlines()
            if line.strip()
        ]

        self.assertEqual(
            len(events),
            4,
        )

        self.assertEqual(
            events[0]["type"],
            "started",
        )

        self.assertEqual(
            events[1]["type"],
            "requirement_started",
        )

        self.assertEqual(
            events[2]["type"],
            "requirement_completed",
        )

        final_event = events[-1]

        self.assertEqual(
            final_event["type"],
            "completed",
        )

        self.assertTrue(
            final_event["ok"]
        )

        self.assertEqual(
            final_event["progress"],
            100,
        )

        self.assertIn(
            "result_id",
            final_event,
        )

        self.assertIn(
            "download_url",
            final_event,
        )

        download_response = self.client.get(
            final_event["download_url"]
        )

        self.assertEqual(
            download_response.status_code,
            200,
        )

        workbook = load_workbook(
            BytesIO(
                download_response.content
            ),
            read_only=True,
        )

        worksheet = workbook[
            "Test Cases"
        ]

        self.assertEqual(
            worksheet["C2"].value,
            "CFC.003.001.001",
        )

        workbook.close()

    def test_rejects_stream_without_assigned_to(
        self,
    ) -> None:
        """Rechaza el stream antes de prepararlo."""
        uploaded_file = SimpleUploadedFile(
            name="document.docx",
            content=b"test-document",
            content_type=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
        )

        response = self.client.post(
            reverse(
                "test_cases:generate_stream"
            ),
            {
                "document": uploaded_file,
                "assigned_to": "",
            },
        )

        self.assertEqual(
            response.status_code,
            400,
        )

        self.assertFalse(
            response.streaming
        )

        self.assertEqual(
            response.json()["code"],
            "ERR_GENERATION_INPUT",
        )

    @patch(
        "apps.test_cases.views."
        "iter_generate_prepared_document"
    )
    @patch(
        "apps.test_cases.views."
        "prepare_generation_document"
    )
    def test_streams_controlled_error_event(
        self,
        prepare_mock: Mock,
        iter_generate_mock: Mock,
    ) -> None:
        """Convierte un error del generador en un evento NDJSON."""
        from apps.test_cases.exceptions import (
            JsonGenerationError,
        )

        prepare_mock.return_value = PreparedGeneration(
            filename="CFC.003_PDD.docx",
            project_id="CFC.003",
            context_text="Contexto",
            blocks=(
                RequirementBlock(
                    requirement_number=1,
                    scenario_name="Validar",
                    input_text="Contenido",
                ),
            ),
        )

        def failing_iterator():
            yield {
                "type": "started",
                "ok": True,
                "progress": 0,
            }

            raise JsonGenerationError(
                "Respuesta JSON inválida."
            )

        iter_generate_mock.return_value = (
            failing_iterator()
        )

        uploaded_file = SimpleUploadedFile(
            name="CFC.003_PDD.docx",
            content=b"document-content",
        )

        response = self.client.post(
            reverse(
                "test_cases:generate_stream"
            ),
            {
                "document": uploaded_file,
                "assigned_to": "Usuario QA",
                "selected_requirements": "1",
            },
        )

        raw_content = b"".join(
            response.streaming_content
        ).decode("utf-8")

        events = [
            json.loads(line)
            for line in raw_content.splitlines()
            if line.strip()
        ]

        self.assertEqual(
            events[-1]["type"],
            "error",
        )

        self.assertFalse(
            events[-1]["ok"]
        )

        self.assertEqual(
            events[-1]["code"],
            "ERR_INVALID_JSON",
        )

    def test_returns_404_without_generated_result(
        self,
    ) -> None:
        """Rechaza una descarga sin resultado previo."""
        response = self.client.get(
            reverse("test_cases:download")
        )

        self.assertEqual(
            response.status_code,
            404,
        )

    @staticmethod
    def _build_valid_docx() -> bytes:
        """Construye un documento Beecker válido."""
        document = Document()

        document.add_paragraph(
            "ID proyecto CFC.003"
        )

        document.add_paragraph(
            "2.4 Acciones detalladas del proceso TO-BE"
        )

        document.add_paragraph(
            "1. Nombre de la acción: Validar archivo"
        )

        document.add_paragraph(
            "Descripción general: Validar estructura."
        )

        document.add_paragraph(
            "2. Nombre de la acción: Procesar información"
        )

        document.add_paragraph(
            "Descripción general: Procesar registros."
        )

        document.add_paragraph(
            "2.5 Matriz de criterios de aceptación"
        )

        output = BytesIO()
        document.save(output)

        return output.getvalue()

    @staticmethod
    def _build_docx_without_requirements() -> bytes:
        """Construye un documento válido sin requerimientos."""
        document = Document()

        document.add_paragraph(
            "Documento sin requerimientos funcionales."
        )

        output = BytesIO()
        document.save(output)

        return output.getvalue()

    @staticmethod
    def _build_generation_result() -> dict[str, object]:
        """Construye un resultado XLSX simulado."""
        row = [""] * ADO_NCOLS

        row[1] = "Test Case"
        row[2] = "CFC.003.001.001"
        row[6] = "Functional"
        row[7] = "1"
        row[8] = "Resultado esperado."
        row[9] = "Procesar correctamente el archivo."
        row[10] = "(Happy Path) - Validar archivo"
        row[12] = "Design"
        row[13] = "CFC.003"
        row[14] = "Usuario QA"

        return {
            "filename": "CFC.003_PDD_TC.xlsx",
            "xlsx_bytes": generate_xlsx(
                [
                    row,
                ]
            ),
            "usage": {
                "input_tokens": 10,
                "output_tokens": 5,
                "total_tokens": 15,
            },
            "cost": {
                "currency": "USD",
                "input_usd": 0,
                "output_usd": 0,
                "total_usd": 0,
                "total_usd_formatted": "$0.000000",
            },
            "elapsed": 1.2,
            "stats": {
                "project_id": "CFC.003",
                "requirements_total": 1,
                "test_cases_total": 1,
                "requirements_not_testable": 0,
            },
            "selected_requirements": [1],
            "missing_selected": [],
        }