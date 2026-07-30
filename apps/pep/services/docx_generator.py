"""
Escritura del documento PEP en formato DOCX mediante placeholders.

Este módulo toma el contexto combinado PAP + PDD/FDD y llena una copia
de la plantilla PEP integrada reemplazando marcas configurables como
**Name_Project, **ID_Project o **Titles_requirements_Tobe.
"""

from __future__ import annotations

from datetime import date
from io import BytesIO
from math import ceil
from typing import Any, Iterable

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph

from apps.pep.services.context_builder import PepContext
from apps.pep.services.template_loader import (
    read_pep_template_bytes,
)

FONT_NAME = "Montserrat Medium"
FONT_SIZE_PT = 10
FONT_COLOR = RGBColor(0x59, 0x59, 0x59)
DEFAULT_TEXT_STYLE = {
    "font_name": "Montserrat Medium",
    "font_size": 10,
    "bold": False,
    "alignment": WD_ALIGN_PARAGRAPH.LEFT,
}

PLACEHOLDER_STYLES = {
    "**Name_Project_T": {
        "font_name": "Open Sans",
        "font_size": 32,
        "bold": True,
        "alignment": WD_ALIGN_PARAGRAPH.CENTER,
    },
    "**ID_Project_H": {
        "font_name": "Montserrat Medium",
        "font_size": 8,
        "bold": False,
        "alignment": WD_ALIGN_PARAGRAPH.CENTER,
    },
    "**Date_issue": {
        "font_name": "Montserrat Medium",
        "font_size": 8,
        "bold": False,
        "alignment": WD_ALIGN_PARAGRAPH.CENTER,
    },
}

SUPPLY_HEADER_FILL = "6A00D4"
SUPPLY_HEADER_TEXT = RGBColor(0xFF, 0xFF, 0xFF)
SUPPLY_BODY_TEXT = RGBColor(0x59, 0x59, 0x59)
SUPPLY_BORDER_COLOR = "7F7F7F"
SUPPLY_FONT_NAME = "Montserrat Medium"
SUPPLY_HEADER_FONT_SIZE = 7
SUPPLY_FONT_SIZE = 8
SUPPLY_NOTE_FONT_SIZE = 8


def generate_pep_docx_bytes(context: PepContext) -> bytes:
    """
    Genera un PEP en DOCX a partir del contexto combinado.

    Args:
        context: Contexto con datos PAP y requerimientos TO-BE.

    Returns:
        Contenido binario del DOCX generado.
    """
    template_bytes = read_pep_template_bytes()
    document = Document(BytesIO(template_bytes))

    replacements = _build_placeholder_replacements(context)
    _replace_placeholders(document, replacements)
    _insert_supply_calculation_table(document, context)

    output = BytesIO()
    document.save(output)

    return output.getvalue()


def _build_placeholder_replacements(
    context: PepContext,
) -> dict[str, str]:
    """
    Construye el mapa de placeholders contra valores finales.

    Args:
        context: Contexto combinado del PEP.

    Returns:
        Diccionario con placeholders y textos de reemplazo.
    """
    pap = context.pap
    roles = pap.roles

    project_name = _safe_text(pap.nombre_proyecto)
    project_id = _safe_text(context.project_id)

    return {
        "**Name_Project_T": f'"{project_name}"',
        "**Name_Project": project_name,

        "**ID_Project_H": project_id,
        "**ID_Project": project_id,

        "**Client_Name": _safe_text(pap.nombre_cliente),
        "**Tecnology_Name": _get_pdd_technology_value(context),
        "**Date_issue": _format_issue_date(date.today()),

        "**Dev_Name": _join_people_names(roles.desarrollador),
        "**Tester_Name": _join_people_names(roles.tester),
        "**SC_Name": _join_people_names(roles.scrum_master),
        "**DM_Name": _join_people_names(roles.delivery_manager),
        "**BA_Name": _join_people_names(roles.business_analyst),
        "**Architec_Name": _join_people_names(roles.arquitecto),
        "**CR_Name": _join_people_names(roles.code_reviewer),

        "**Software_requirements": _build_bullet_list(
            pap.requisitos_software.items,
        ),
        "**Hardware_Requirements": _build_bullet_list(
            pap.requisitos_hardware.items,
        ),
        "**Titles_requirements_Tobe": _build_pdd_requirement_lines(context),
    }


def _insert_supply_calculation_table(
    document: Any,
    context: PepContext,
) -> None:
    """
    Inserta la tabla de cálculo de insumos en el placeholder indicado.

    Args:
        document: Documento DOCX cargado.
        context: Contexto combinado del PEP.
    """
    placeholder = "**Supply_Calculation"

    target_paragraph = next(
        (
            paragraph
            for paragraph in _iter_all_paragraphs(document)
            if placeholder in paragraph.text
        ),
        None,
    )

    if target_paragraph is None:
        return

    _clear_paragraph_runs(target_paragraph)

    plan = context.pdd.calculo_insumos.plan_insumos

    if plan is None:
        target_paragraph.add_run(
            "No se pudo calcular la tabla de insumos."
        )
        _apply_pep_run_style(
            target_paragraph.runs[0],
            DEFAULT_TEXT_STYLE,
        )
        return

    table = document.add_table(
        rows=1,
        cols=6,
    )

    _apply_safe_table_grid_style(table)

    header_cells = table.rows[0].cells
    headers = [
        (
            "Fase del Proceso\n"
            "(Fase de las pruebas y nivel de prueba)"
        ),
        "Cantidad de insumos",
        (
            "% de Representación\n"
            "en la Ejecución"
        ),
        (
            "Frecuencia\n"
            "(Ejecución en entornos reales)"
        ),
        "Tipo de dato\n(Tipo de insumo)",
        (
            "Característica de los insumos\n"
            "(Acorde al contexto)"
        ),
    ]

    for index, header in enumerate(headers):
        header_cells[index].text = header

    for phase in plan.fases_prueba:
        row_cells = table.add_row().cells

        row_values = [
            (
                f"{phase.fase_proceso} "
                f"({phase.nivel_prueba})"
            ),
            _format_quantity(
                phase.cantidad,
                phase.unidad_elemento,
            ),
            f"{phase.porcentaje_aplicado}%",
            _format_optional_table_text(
                phase.frecuencia,
            ),
            _format_optional_table_text(
                phase.tipo_dato,
            ),
            _format_characteristics(
                phase.caracteristicas,
            ),
        ]

        for index, value in enumerate(row_values):
            row_cells[index].text = value

    _style_supply_calculation_table(table)

    target_paragraph._p.addnext(table._tbl)

    note_parts: list[str] = []

    calculation_criterion = (
        plan.criterio_calculo or ""
    ).strip()

    deployment_note = (
        plan.nota_deployment or ""
    ).strip()

    if calculation_criterion:
        note_parts.append(
            "Criterio de cálculo: "
            f"{calculation_criterion}"
        )

    if deployment_note:
        note_parts.append(
            "Deployment/UAT: "
            f"{deployment_note}"
        )

    note_text = "\n".join(note_parts)

    if note_text:
        note_paragraph = document.add_paragraph()
        note_paragraph.paragraph_format.space_before = Pt(4)
        note_paragraph.paragraph_format.space_after = Pt(0)

        note_run = note_paragraph.add_run(
            note_text,
        )

        note_run.font.name = SUPPLY_FONT_NAME
        note_run.font.size = Pt(
            SUPPLY_NOTE_FONT_SIZE
        )
        note_run.font.color.rgb = SUPPLY_BODY_TEXT

        _set_run_font_family(
            note_run,
            SUPPLY_FONT_NAME,
        )

        table._tbl.addnext(
            note_paragraph._p
        )

def _apply_safe_table_grid_style(table) -> None:
    """
    Aplica un estilo de tabla con cuadrícula sin depender de que la
    plantilla tenga registrado el estilo 'Table Grid'.

    Args:
        table: Tabla de python-docx.
    """
    possible_style_names = (
        "Table Grid",
        "Tabla con cuadrícula",
        "Cuadrícula de tabla",
    )

    for style_name in possible_style_names:
        try:
            table.style = style_name
            return
        except KeyError:
            continue

    _apply_table_borders(table)


def _apply_table_borders(
    table: Any,
    *,
    color: str = "808080",
) -> None:
    """
    Aplica bordes visibles a una tabla mediante XML.

    Args:
        table: Tabla de python-docx.
        color: Color hexadecimal sin #.
    """
    table_properties = table._tbl.tblPr

    borders = table_properties.first_child_found_in(
        "w:tblBorders",
    )

    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table_properties.append(borders)

    for border_name in (
        "top",
        "left",
        "bottom",
        "right",
        "insideH",
        "insideV",
    ):
        border = borders.find(
            qn(f"w:{border_name}"),
        )

        if border is None:
            border = OxmlElement(
                f"w:{border_name}",
            )
            borders.append(border)

        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "5")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)


def _format_cell_text(
    cell: Any,
    *,
    bold: bool,
) -> None:
    """
    Aplica formato base a una celda de la tabla de insumos.
    """
    style = dict(DEFAULT_TEXT_STYLE)
    style["bold"] = bold

    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            _apply_pep_run_style(
                run,
                style,
            )


def _style_supply_calculation_table(table: Any) -> None:
    """
    Aplica formato visual a la tabla de cálculo de insumos.

    Args:
        table: Tabla creada con python-docx.
    """
    table.autofit = False

    column_widths = (
        Inches(1.45),
        Inches(1.00),
        Inches(1.10),
        Inches(1.20),
        Inches(1.10),
        Inches(1.70),
    )

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = column_widths[index]
            cell.vertical_alignment = (
                WD_CELL_VERTICAL_ALIGNMENT.CENTER
            )
            _set_cell_margins(
                cell,
                top=80,
                start=80,
                bottom=80,
                end=80,
            )

    _apply_table_borders(
        table,
        color=SUPPLY_BORDER_COLOR,
    )

    for cell in table.rows[0].cells:
        _set_cell_background(
            cell,
            SUPPLY_HEADER_FILL,
        )

        _set_cell_text_style(
            cell=cell,
            bold=True,
            color=SUPPLY_HEADER_TEXT,
            font_size=SUPPLY_HEADER_FONT_SIZE,
        )

    for row in table.rows[1:]:
        for cell in row.cells:
            _set_cell_text_style(
                cell=cell,
                bold=False,
                color=SUPPLY_BODY_TEXT,
                font_size=SUPPLY_FONT_SIZE,
            )

def _set_cell_background(
    cell: Any,
    fill: str,
) -> None:
    """
    Aplica color de fondo a una celda.

    Args:
        cell: Celda de python-docx.
        fill: Color hexadecimal sin #.
    """
    cell_properties = cell._tc.get_or_add_tcPr()

    shading = cell_properties.find(
        qn("w:shd"),
    )

    if shading is None:
        shading = OxmlElement("w:shd")
        cell_properties.append(shading)

    shading.set(
        qn("w:fill"),
        fill,
    )


def _set_cell_text_style(
    *,
    cell: Any,
    bold: bool,
    color: RGBColor,
    font_size: float,
) -> None:
    """
    Aplica estilo de texto a todos los runs de una celda.

    Args:
        cell: Celda de python-docx.
        bold: Indica si el texto será negrita.
        color: Color RGB del texto.
        font_size: Tamaño de fuente.
    """
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)

        for run in paragraph.runs:
            run.bold = bold
            run.font.name = SUPPLY_FONT_NAME
            run.font.size = Pt(font_size)
            run.font.color.rgb = color
            _set_run_font_family(
                run,
                SUPPLY_FONT_NAME,
            )


def _set_cell_margins(
    cell: Any,
    *,
    top: int,
    start: int,
    bottom: int,
    end: int,
) -> None:
    """
    Ajusta márgenes internos de una celda.

    Args:
        cell: Celda de python-docx.
        top: Margen superior en twentieths of a point.
        start: Margen izquierdo.
        bottom: Margen inferior.
        end: Margen derecho.
    """
    cell_properties = cell._tc.get_or_add_tcPr()

    margins = cell_properties.first_child_found_in(
        "w:tcMar",
    )

    if margins is None:
        margins = OxmlElement("w:tcMar")
        cell_properties.append(margins)

    margin_values = {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }

    for margin_name, margin_value in margin_values.items():
        margin = margins.find(
            qn(f"w:{margin_name}"),
        )

        if margin is None:
            margin = OxmlElement(
                f"w:{margin_name}",
            )
            margins.append(margin)

        margin.set(
            qn("w:w"),
            str(margin_value),
        )

        margin.set(
            qn("w:type"),
            "dxa",
        )


def _set_run_font_family(
    run: Any,
    font_name: str,
) -> None:
    """
    Aplica la familia tipográfica al run, incluyendo eastAsia.

    Args:
        run: Run DOCX.
        font_name: Nombre de la fuente.
    """
    run_properties = run._element.get_or_add_rPr()

    fonts = run_properties.rFonts

    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        run_properties.append(fonts)

    fonts.set(
        qn("w:ascii"),
        font_name,
    )
    fonts.set(
        qn("w:hAnsi"),
        font_name,
    )
    fonts.set(
        qn("w:eastAsia"),
        font_name,
    )


def _format_quantity(
    quantity: float | int | None,
    unit: str | None,
) -> str:
    """
    Formatea una cantidad con separadores y unidad.
    """
    if quantity is None:
        return "N/A"

    formatted_quantity = f"{ceil(quantity):,}"

    if not unit:
        return formatted_quantity

    return f"{formatted_quantity} {unit}"


def _format_optional_table_text(
    value: str | None,
) -> str:
    """
    Formatea un valor opcional para una celda.
    """
    clean_value = " ".join(
        (value or "").split()
    )

    return clean_value or "N/A"


def _format_characteristics(
    characteristics: list[str],
) -> str:
    """
    Formatea las características como lista multilínea.
    """
    clean_characteristics = [
        clean_value
        for value in characteristics
        if (
            clean_value := " ".join(
                (value or "").split()
            )
        )
    ]

    if not clean_characteristics:
        return "N/A"

    return "\n".join(
        f"• {value}"
        for value in clean_characteristics
    )

def _get_pdd_technology_value(
    context: PepContext,
) -> str:
    """
    Obtiene la tecnología detectada desde el PDD/FDD.
    """
    value = context.pdd.tecnologia.valor

    return value.strip() if value else ""


def _replace_placeholders(
    document: Any,
    replacements: dict[str, str],
) -> None:
    """
    Reemplaza placeholders en todos los párrafos del documento.

    Incluye párrafos normales, tablas, encabezados y pies de página.

    Args:
        document: Documento DOCX cargado.
        replacements: Diccionario placeholder -> valor.
    """
    for paragraph in _iter_all_paragraphs(document):
        _replace_placeholders_in_paragraph(paragraph, replacements)


def _replace_placeholders_in_paragraph(
    paragraph: Any,
    replacements: dict[str, str],
) -> None:
    """
    Reemplaza placeholders dentro de un párrafo.

    Args:
        paragraph: Párrafo DOCX.
        replacements: Diccionario placeholder -> valor.
    """
    original_text = paragraph.text
    if not original_text:
        return

    matched_placeholders = [
        placeholder
        for placeholder in sorted(replacements, key=len, reverse=True)
        if placeholder in original_text
    ]

    if not matched_placeholders:
        return

    new_text = original_text

    for placeholder in matched_placeholders:
        new_text = new_text.replace(
            placeholder,
            replacements[placeholder],
        )

    style = _resolve_style_for_placeholders(matched_placeholders)

    _set_paragraph_text(
        paragraph,
        new_text,
        style,
    )

def _resolve_style_for_placeholders(
    placeholders: list[str],
) -> dict[str, Any]:
    """
    Resuelve el estilo a aplicar según el placeholder encontrado.

    Args:
        placeholders: Lista de placeholders encontrados.

    Returns:
        Diccionario de estilo.
    """
    style = dict(DEFAULT_TEXT_STYLE)

    for placeholder in placeholders:
        custom_style = PLACEHOLDER_STYLES.get(placeholder)
        if custom_style:
            style.update(custom_style)
            return style

    return style

def _iter_all_paragraphs(document: Any) -> Iterable[Any]:
    """
    Itera todos los párrafos del documento, incluyendo tablas,
    encabezados, pies de página y cuadros de texto.

    Args:
        document: Documento DOCX.

    Yields:
        Párrafos encontrados.
    """
    yield from document.paragraphs
    yield from _iter_textbox_paragraphs(document)

    for table in document.tables:
        yield from _iter_table_paragraphs(table)

    for section in document.sections:
        yield from section.header.paragraphs
        yield from _iter_textbox_paragraphs(section.header)

        yield from section.footer.paragraphs
        yield from _iter_textbox_paragraphs(section.footer)

        for table in section.header.tables:
            yield from _iter_table_paragraphs(table)

        for table in section.footer.tables:
            yield from _iter_table_paragraphs(table)

def _iter_textbox_paragraphs(parent: Any) -> Iterable[Any]:
    """
    Itera párrafos dentro de cuadros de texto.

    Args:
        parent: Documento, encabezado o pie de página.

    Yields:
        Párrafos encontrados dentro de text boxes.
    """
    for paragraph_element in parent._element.xpath(".//w:txbxContent//w:p"):
        yield Paragraph(paragraph_element, parent)

def _iter_table_paragraphs(table: Any) -> Iterable[Any]:
    """
    Itera párrafos dentro de una tabla, incluyendo tablas anidadas.

    Args:
        table: Tabla DOCX.

    Yields:
        Párrafos encontrados dentro de las celdas.
    """
    for row in table.rows:
        for cell in row.cells:
            yield from cell.paragraphs

            for nested_table in cell.tables:
                yield from _iter_table_paragraphs(nested_table)


def _set_paragraph_text(
    paragraph: Any,
    text: str,
    style: dict[str, Any],
) -> None:
    """
    Reemplaza el texto de un párrafo aplicando formato configurable.

    Args:
        paragraph: Párrafo DOCX.
        text: Texto a escribir.
        style: Estilo a aplicar.
    """
    _clear_paragraph_runs(paragraph)

    paragraph.alignment = style["alignment"]

    lines = text.splitlines() or [""]

    first_run = paragraph.add_run(lines[0])
    _apply_pep_run_style(first_run, style)

    for line in lines[1:]:
        run = paragraph.add_run()
        run.add_break()
        run.add_text(line)
        _apply_pep_run_style(run, style)


def _clear_paragraph_runs(paragraph: Any) -> None:
    """
    Elimina los runs existentes de un párrafo.

    Args:
        paragraph: Párrafo DOCX.
    """
    for run in list(paragraph.runs):
        run_element = run._element
        run_element.getparent().remove(run_element)


def _apply_pep_run_style(
    run: Any,
    style: dict[str, Any],
) -> None:
    """
    Aplica formato configurable para campos editados del PEP.

    Args:
        run: Run DOCX.
        style: Diccionario de estilo.
    """
    font_name = style["font_name"]

    run.font.name = font_name
    run.font.size = Pt(style["font_size"])
    run.font.bold = style["bold"]
    run.font.color.rgb = FONT_COLOR

    _set_run_font_family(
        run,
        font_name,
    )


def _join_people_names(names: list[str] | None) -> str:
    """
    Convierte una lista de personas en texto multilinea.

    Args:
        names: Lista de nombres detectados.

    Returns:
        Texto con una persona por línea o N/A.
    """
    cleaned_names = [
        name.strip()
        for name in (names or [])
        if (name or "").strip()
    ]

    if not cleaned_names:
        return "N/A"

    return "\n".join(cleaned_names)


def _build_bullet_list(items: list[str]) -> str:
    """
    Construye una lista con viñetas para insertar en el PEP.

    Args:
        items: Elementos extraídos del PAP.

    Returns:
        Texto con viñetas o N/A.
    """
    cleaned_items = [
        item.strip()
        for item in items
        if (item or "").strip()
    ]

    if not cleaned_items:
        return "N/A"

    return "\n".join(f"• {item}" for item in cleaned_items)


def _build_pdd_requirement_lines(
    context: PepContext,
) -> str:
    """
    Construye la lista de requerimientos funcionales extraídos del PDD/FDD.

    Los títulos son utilizados exactamente como fueron validados en el
    análisis del documento, sin agregar numeración adicional.

    Args:
        context: Contexto combinado del PEP.

    Returns:
        Texto con un requerimiento funcional por línea.
    """
    requirements = [
        requirement.strip()
        for requirement in context.pdd.requerimientos
        if requirement.strip()
    ]

    if not requirements:
        return "No se detectaron requerimientos funcionales."

    return "\n".join(requirements)

def _format_issue_date(current_date: date) -> str:
    """
    Formatea la fecha de emisión.

    Args:
        current_date: Fecha actual.

    Returns:
        Fecha en formato dd de mm de aaaa.
    """
    return (
        f"{current_date.day:02d} de "
        f"{current_date.month:02d} de "
        f"{current_date.year}"
    )

def _safe_text(value: str | None) -> str:
    """
    Convierte valores nulos o vacíos a N/A.

    Args:
        value: Texto opcional.

    Returns:
        Texto seguro.
    """
    clean = (value or "").strip()
    return clean or "N/A"