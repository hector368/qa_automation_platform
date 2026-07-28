from pathlib import Path
from tempfile import TemporaryDirectory

from django.test import SimpleTestCase
from docx import Document

from apps.pep.exceptions import PepTemplateError
from apps.pep.services.template_loader import (
    get_pep_template_path,
    read_pep_template_bytes,
    validate_pep_template,
)


class PepTemplateLoaderTests(SimpleTestCase):
    def test_integrated_template_exists(self) -> None:
        path = get_pep_template_path()

        self.assertTrue(
            path.is_file(),
        )

        self.assertEqual(
            path.name,
            "pep_template.docx",
        )

        self.assertIn(
            "pep",
            path.parts,
        )

    def test_validates_docx_template(self) -> None:
        with TemporaryDirectory() as directory:
            path = Path(directory) / "template.docx"

            document = Document()
            document.add_paragraph("PEP")
            document.save(path)

            result = validate_pep_template(
                path
            )

            self.assertEqual(
                result.filename,
                "template.docx",
            )

            self.assertGreater(
                result.size_bytes,
                0,
            )

    def test_reads_template_bytes(self) -> None:
        with TemporaryDirectory() as directory:
            path = Path(directory) / "template.docx"

            document = Document()
            document.add_paragraph("PEP")
            document.save(path)

            content = read_pep_template_bytes(
                path
            )

            self.assertGreater(
                len(content),
                0,
            )

    def test_rejects_missing_template(self) -> None:
        with self.assertRaises(
            PepTemplateError,
        ):
            validate_pep_template(
                Path("template_inexistente.docx")
            )

    def test_rejects_non_docx_template(self) -> None:
        with TemporaryDirectory() as directory:
            path = Path(directory) / "template.txt"

            path.write_text(
                "contenido",
                encoding="utf-8",
            )

            with self.assertRaises(
                PepTemplateError,
            ):
                validate_pep_template(
                    path
                )