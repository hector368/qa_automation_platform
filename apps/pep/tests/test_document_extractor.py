from io import BytesIO

import fitz
from django.test import SimpleTestCase
from docx import Document

from apps.pep.exceptions import (
    DocumentExtractionError,
    EmptyDocumentTextError,
    EmptyFileError,
)
from apps.pep.services.document_extractor import (
    extract_text_from_document,
)


class DocumentExtractorTests(SimpleTestCase):
    def test_extracts_pdf_text(self) -> None:
        pdf_bytes = self._build_pdf(
            "ID proyecto CFC.003"
        )

        text = extract_text_from_document(
            filename="sample.pdf",
            file_bytes=pdf_bytes,
        )

        self.assertIn(
            "ID proyecto CFC.003",
            text,
        )

    def test_extracts_docx_paragraphs_and_tables(
        self,
    ) -> None:
        docx_bytes = self._build_docx()

        text = extract_text_from_document(
            filename="sample.docx",
            file_bytes=docx_bytes,
        )

        self.assertIn(
            "Acciones detalladas",
            text,
        )

        self.assertIn(
            "Nombre de la acción",
            text,
        )

    def test_rejects_empty_bytes(self) -> None:
        with self.assertRaises(
            EmptyFileError,
        ):
            extract_text_from_document(
                filename="sample.pdf",
                file_bytes=b"",
            )

    def test_rejects_corrupt_pdf(self) -> None:
        with self.assertRaises(
            DocumentExtractionError,
        ):
            extract_text_from_document(
                filename="sample.pdf",
                file_bytes=b"not-a-pdf",
            )

    def test_rejects_pdf_without_text(self) -> None:
        empty_pdf = self._build_pdf("")

        with self.assertRaises(
            EmptyDocumentTextError,
        ):
            extract_text_from_document(
                filename="empty.pdf",
                file_bytes=empty_pdf,
            )

    @staticmethod
    def _build_pdf(text: str) -> bytes:
        document = fitz.open()
        page = document.new_page()

        if text:
            page.insert_text(
                (72, 72),
                text,
            )

        content = document.tobytes()
        document.close()

        return content

    @staticmethod
    def _build_docx() -> bytes:
        document = Document()
        document.add_paragraph(
            "2.4 Acciones detalladas"
        )

        table = document.add_table(
            rows=1,
            cols=2,
        )

        table.cell(0, 0).text = "1"
        table.cell(0, 1).text = (
            "Nombre de la acción"
        )

        output = BytesIO()
        document.save(output)

        return output.getvalue()
