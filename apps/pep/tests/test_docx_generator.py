from io import BytesIO
from unittest.mock import patch

from django.test import SimpleTestCase
from docx import Document

from apps.pep.services.context_builder import (
    build_pep_context,
)
from apps.pep.services.docx_generator import (
    generate_pep_docx_bytes,
)
from apps.pep.tests.pep_factories import (
    build_pap_data,
    build_pdd_data,
)


class PepDocxGeneratorTests(SimpleTestCase):
    @patch(
        "apps.pep.services.docx_generator."
        "read_pep_template_bytes"
    )
    def test_generates_document_and_replaces_markers(
        self,
        template_mock,
    ) -> None:
        template_mock.return_value = (
            self._build_template_bytes()
        )

        context = build_pep_context(
            pap_data=build_pap_data(),
            pdd_data=build_pdd_data(),
        )

        output = generate_pep_docx_bytes(
            context
        )

        self.assertGreater(
            len(output),
            0,
        )

        document = Document(
            BytesIO(output)
        )

        document_text = self._extract_document_text(
            document
        )

        # Datos generales del proyecto.
        self.assertIn(
            "Proyecto de prueba",
            document_text,
        )

        self.assertIn(
            "CFC.003",
            document_text,
        )

        self.assertIn(
            "Cliente de prueba",
            document_text,
        )

        # La tecnología debe provenir del PDD/FDD.
        self.assertIn(
            "Power Automate Desktop",
            document_text,
        )

        self.assertNotIn(
            "UiPath",
            document_text,
        )

        # Requerimientos funcionales.
        self.assertIn(
            "Consultar solicitudes",
            document_text,
        )

        self.assertIn(
            "Validar información",
            document_text,
        )

        # Encabezados de la nueva tabla.
        self.assertIn(
            "Fase del Proceso",
            document_text,
        )

        self.assertIn(
            "Cantidad de insumos",
            document_text,
        )

        self.assertIn(
            "% de Representación",
            document_text,
        )

        self.assertIn(
            "Frecuencia",
            document_text,
        )

        self.assertIn(
            "Tipo de dato",
            document_text,
        )

        self.assertIn(
            "Característica de los insumos",
            document_text,
        )

        # Fases y niveles de prueba.
        self.assertIn(
            "Planificación",
            document_text,
        )

        self.assertIn(
            "Pruebas Unitarias",
            document_text,
        )

        self.assertIn(
            "Preparación",
            document_text,
        )

        self.assertIn(
            "Pruebas de Integración",
            document_text,
        )

        self.assertIn(
            "Ejecución",
            document_text,
        )

        self.assertIn(
            "Pruebas de Sistema / End-to-End",
            document_text,
        )

        self.assertIn(
            "Cierre",
            document_text,
        )

        self.assertIn(
            "Pruebas de Aceptación / UAT",
            document_text,
        )

        # Se conservan los criterios anteriores.
        self.assertIn(
            "50%",
            document_text,
        )

        self.assertIn(
            "120%",
            document_text,
        )

        # Contexto documental de los insumos.
        self.assertIn(
            "Excel / ServiceNow",
            document_text,
        )

        self.assertIn(
            "Datos transaccionales",
            document_text,
        )

        # Notas del cálculo.
        self.assertIn(
            "Criterio de cálculo:",
            document_text,
        )

        self.assertIn(
            "Deployment/UAT:",
            document_text,
        )

        self.assertIn(
            "insumos productivos",
            document_text,
        )

        self.assertIn(
            "entorno productivo",
            document_text,
        )

        # No deben quedar marcadores pendientes.
        self.assertNotIn(
            "**",
            document_text,
        )
    @staticmethod
    def _build_template_bytes() -> bytes:
        document = Document()

        document.add_paragraph(
            "**Name_Project_T"
        )

        document.add_paragraph(
            "Nombre: **Name_Project"
        )

        document.add_paragraph(
            "ID: **ID_Project"
        )

        document.add_paragraph(
            "Cliente: **Client_Name"
        )

        document.add_paragraph(
            "Tecnología: **Tecnology_Name"
        )

        document.add_paragraph(
            "**Software_requirements"
        )

        document.add_paragraph(
            "**Hardware_Requirements"
        )

        document.add_paragraph(
            "**Titles_requirements_Tobe"
        )

        document.add_paragraph(
            "**Supply_Calculation"
        )

        table = document.add_table(
            rows=2,
            cols=4,
        )

        table.cell(0, 0).text = "**Dev_Name"
        table.cell(0, 1).text = "**Tester_Name"
        table.cell(0, 2).text = "**SC_Name"
        table.cell(0, 3).text = "**DM_Name"

        table.cell(1, 0).text = "**BA_Name"
        table.cell(1, 1).text = "**Architec_Name"
        table.cell(1, 2).text = "**CR_Name"
        table.cell(1, 3).text = "**Date_issue"

        header = document.sections[0].header

        header.paragraphs[0].text = (
            "Proyecto **ID_Project_H"
        )

        output = BytesIO()
        document.save(output)

        return output.getvalue()

    @classmethod
    def _extract_document_text(
        cls,
        document: Document,
    ) -> str:
        parts = [
            paragraph.text
            for paragraph in document.paragraphs
        ]

        for table in document.tables:
            parts.extend(
                cls._extract_table_text(
                    table
                )
            )

        for section in document.sections:
            parts.extend(
                paragraph.text
                for paragraph
                in section.header.paragraphs
            )

            parts.extend(
                paragraph.text
                for paragraph
                in section.footer.paragraphs
            )

        return "\n".join(parts)

    @classmethod
    def _extract_table_text(
        cls,
        table,
    ) -> list[str]:
        parts: list[str] = []

        for row in table.rows:
            for cell in row.cells:
                parts.extend(
                    paragraph.text
                    for paragraph in cell.paragraphs
                )

                for nested_table in cell.tables:
                    parts.extend(
                        cls._extract_table_text(
                            nested_table
                        )
                    )

        return parts
    
    def test_integrated_template_contains_required_markers(
        self,
    ) -> None:
        from apps.pep.services.docx_generator import (
            _iter_all_paragraphs,
        )
        from apps.pep.services.template_loader import (
            read_pep_template_bytes,
        )

        template = Document(
            BytesIO(
                read_pep_template_bytes()
            )
        )

        template_text = "\n".join(
            paragraph.text
            for paragraph in _iter_all_paragraphs(
                template
            )
        )

        required_markers = {
            "**Name_Project_T",
            "**Name_Project",
            "**ID_Project_H",
            "**ID_Project",
            "**Client_Name",
            "**Tecnology_Name",
            "**Date_issue",
            "**Dev_Name",
            "**Tester_Name",
            "**SC_Name",
            "**DM_Name",
            "**BA_Name",
            "**Architec_Name",
            "**CR_Name",
            "**Software_requirements",
            "**Hardware_Requirements",
            "**Titles_requirements_Tobe",
            "**Supply_Calculation",
        }

        missing_markers = sorted(
            marker
            for marker in required_markers
            if marker not in template_text
        )

        self.assertEqual(
            missing_markers,
            [],
            msg=(
                "La plantilla no contiene estos "
                f"marcadores: {missing_markers}"
            ),
        )